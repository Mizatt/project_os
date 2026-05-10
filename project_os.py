"""
RPM Injeção Eletrônica Diesel — Gerador de Ordem de Serviço
Versão GUI (tkinter): formulário visual com botão Gerar O.S.

Dependências:
    pip install python-docx

Execução:
    python gerar_os_gui.py
"""

import os
import re
import sys
import traceback
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from datetime import datetime

# ─────────────────────────────────────────────
# Log de erros — grava erro_log.txt na pasta do exe
# Útil para diagnosticar falhas ao rodar como .exe
# ─────────────────────────────────────────────
def _pasta_exe():
    """Retorna a pasta onde o executável (ou script) está."""
    if hasattr(sys, "_MEIPASS"):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__))

_log_path = os.path.join(_pasta_exe(), "erro_log.txt")

def _salvar_erro(tipo, valor, tb):
    try:
        with open(_log_path, "w", encoding="utf-8") as f:
            traceback.print_exception(tipo, valor, tb, file=f)
    except Exception:
        pass
    sys.__excepthook__(tipo, valor, tb)

sys.excepthook = _salvar_erro

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─────────────────────────────────────────────
# Constantes
# ─────────────────────────────────────────────
LOGO_FILENAME = "LOGO_Horizontal_cor.png"
EMPRESA_ENDE  = "RODOVIA RAPOSO TAVARES KM 104,750 - 245 - IPANEMA DO MEIO - SOROCABA - SP"
COR_ESCURA    = "17233F"
COR_MEDIA     = "4A90E2"
COR_LINHA_PAR = "EBEBEB"

# Paleta da GUI
GUI_BG        = "#F4F6F9"
GUI_CARD      = "#FFFFFF"
GUI_DARK      = "#17233F"
GUI_ACCENT    = "#4A90E2"
GUI_LABEL     = "#5A6A7A"
GUI_BORDER    = "#D0D8E4"
GUI_BTN_FG    = "#FFFFFF"
GUI_ERROR     = "#D93025"
GUI_SUCCESS   = "#2E7D32"
FONT_TITLE    = ("Segoe UI", 15, "bold")
FONT_SECTION  = ("Segoe UI", 10, "bold")
FONT_LABEL    = ("Segoe UI", 9)
FONT_ENTRY    = ("Segoe UI", 10)
FONT_BTN      = ("Segoe UI", 11, "bold")
FONT_SMALL    = ("Segoe UI", 8)


def _resource_path(filename):
    """Resolve caminhos dentro do bundle PyInstaller ou no diretório normal."""
    if hasattr(sys, "_MEIPASS"):
        return os.path.join(sys._MEIPASS, filename)
    return os.path.join(os.path.dirname(os.path.abspath(__file__)), filename)

def _localizar_logo():
    candidatos = [
        _resource_path(LOGO_FILENAME),                              # bundled / ao lado do script
        os.path.join(os.path.dirname(sys.executable), LOGO_FILENAME),  # ao lado do .exe
        os.path.join(os.getcwd(), LOGO_FILENAME),                   # diretório de trabalho
    ]
    for p in candidatos:
        if os.path.isfile(p):
            return p
    return None

LOGO_PATH = _localizar_logo()


# ─────────────────────────────────────────────
# Helpers docx (idênticos ao terminal)
# ─────────────────────────────────────────────

def _bg(cell, cor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), cor)
    tcPr.append(shd)

def _bordas(cell, cor="AAAAAA", esp="4", top=True, bottom=True, left=True, right=True):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for lado, ativo in (("top",top),("bottom",bottom),("left",left),("right",right)):
        el = OxmlElement("w:"+lado)
        if ativo:
            el.set(qn("w:val"),"single"); el.set(qn("w:sz"),esp); el.set(qn("w:color"),cor)
        else:
            el.set(qn("w:val"),"nil")
        borders.append(el)
    tcPr.append(borders)

def _texto_celula(cell, texto, bold=False, size=10,
                  align=WD_ALIGN_PARAGRAPH.LEFT, cor_hex=None):
    p = cell.paragraphs[0]; p.alignment = align
    run = p.add_run(texto); run.bold = bold; run.font.size = Pt(size)
    if cor_hex: run.font.color.rgb = RGBColor(*bytes.fromhex(cor_hex))
    return run

def _par(doc, texto="", bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, depois=0):
    p = doc.add_paragraph(); p.alignment = align
    p.paragraph_format.space_after = Pt(depois)
    run = p.add_run(texto); run.bold = bold; run.font.size = Pt(size)
    return p

def gerar_os(dados, destino):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin=Cm(1.5); sec.bottom_margin=Cm(1.5)
        sec.left_margin=Cm(2.0); sec.right_margin=Cm(2.0)

    # Cabecalho
    tbl_hdr = doc.add_table(rows=2, cols=2)
    tbl_hdr.style="Table Grid"; tbl_hdr.autofit=False
    tbl_hdr.columns[0].width=Cm(8); tbl_hdr.columns[1].width=Cm(9)
    c_logo = tbl_hdr.cell(0,0)
    _bg(c_logo,"FFFFFF"); _bordas(c_logo,cor="CCCCCC")
    c_logo.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_logo = c_logo.paragraphs[0]; p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO_PATH:
        p_logo.add_run().add_picture(LOGO_PATH, width=Cm(7))
    else:
        r = p_logo.add_run("RPM ELETRODIESEL"); r.bold=True; r.font.size=Pt(14)
        r.font.color.rgb = RGBColor(0x17,0x23,0x3F)
    c_tit = tbl_hdr.cell(0,1)
    _bg(c_tit,"FFFFFF"); _bordas(c_tit,cor="CCCCCC")
    c_tit.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_tit = c_tit.paragraphs[0]; p_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p_tit.add_run("ORDEM DE SERVICO")
    r1.bold=True; r1.font.size=Pt(15); r1.font.color.rgb=RGBColor(0x00,0x00,0x00)
    c_end = tbl_hdr.cell(1,0); c_end.merge(tbl_hdr.cell(1,1))
    _bg(c_end,COR_MEDIA); _bordas(c_end,cor="FFFFFF")
    _texto_celula(c_end, EMPRESA_ENDE, size=9,
                  align=WD_ALIGN_PARAGRAPH.CENTER, cor_hex="FFFFFF")
    doc.add_paragraph()

    # OS / Data
    tbl_os = doc.add_table(rows=1, cols=2)
    tbl_os.style="Table Grid"; tbl_os.autofit=False
    tbl_os.columns[0].width=Cm(8.5); tbl_os.columns[1].width=Cm(8.5)
    for cel, txt in ((tbl_os.cell(0,0),"O.S.: "+dados["os_numero"]),
                     (tbl_os.cell(0,1),"DATA: "+dados["data_os"])):
        _bg(cel,COR_ESCURA); _bordas(cel,cor="FFFFFF")
        aln = WD_ALIGN_PARAGRAPH.RIGHT if "DATA" in txt else WD_ALIGN_PARAGRAPH.LEFT
        _texto_celula(cel,txt,bold=True,size=11,align=aln,cor_hex="FFFFFF")
    doc.add_paragraph()

    # Cliente
    campos_cli = [("CLIENTE",dados["cliente"]),("CIDADE",dados["cidade"]),
                  ("MOTOR",dados["motor"]),("VEICULO",dados["veiculo"]),("PLACA",dados["placa"])]
    tbl_cli = doc.add_table(rows=len(campos_cli), cols=2)
    tbl_cli.style="Table Grid"; tbl_cli.autofit=False
    tbl_cli.columns[0].width=Cm(3.5); tbl_cli.columns[1].width=Cm(13.5)
    for i,(label,valor) in enumerate(campos_cli):
        cl=tbl_cli.cell(i,0); cv=tbl_cli.cell(i,1)
        _bg(cl,COR_MEDIA); _bordas(cl,cor="CCCCCC"); _bordas(cv,cor="CCCCCC")
        _texto_celula(cl,label,bold=True,size=9,cor_hex="FFFFFF")
        _texto_celula(cv,valor,size=10)
    doc.add_paragraph()

    # Servicos
    p = _par(doc,"SERVICOS REALIZADOS",bold=True,size=10,depois=2)
    p.runs[0].font.color.rgb = RGBColor(*bytes.fromhex(COR_ESCURA))
    tbl_svc = doc.add_table(rows=len(dados["servicos"]), cols=1)
    tbl_svc.style="Table Grid"; tbl_svc.autofit=False; tbl_svc.columns[0].width=Cm(17)
    for i,svc in enumerate(dados["servicos"]):
        c=tbl_svc.cell(i,0); _bg(c,COR_LINHA_PAR if i%2==0 else "FFFFFF")
        _bordas(c,cor="CCCCCC"); _texto_celula(c,"{:02d}. {}".format(i+1,svc),size=10)
    doc.add_paragraph()

    # Pecas
    p = _par(doc,"PECAS UTILIZADAS",bold=True,size=10,depois=2)
    p.runs[0].font.color.rgb = RGBColor(*bytes.fromhex(COR_ESCURA))
    tbl_pcs = doc.add_table(rows=1+len(dados["pecas"]), cols=3)
    tbl_pcs.style="Table Grid"; tbl_pcs.autofit=False
    tbl_pcs.columns[0].width=Cm(3); tbl_pcs.columns[1].width=Cm(5); tbl_pcs.columns[2].width=Cm(9)
    for j,titulo in enumerate(("QTDE.","CODIGO","DESCRICAO")):
        ch=tbl_pcs.cell(0,j); _bg(ch,COR_MEDIA); _bordas(ch,cor="FFFFFF")
        _texto_celula(ch,titulo,bold=True,size=9,align=WD_ALIGN_PARAGRAPH.CENTER,cor_hex="FFFFFF")
    for i,peca in enumerate(dados["pecas"]):
        bg = COR_LINHA_PAR if i%2==0 else "FFFFFF"
        for j,(val,aln) in enumerate([(str(peca["qtd"]),WD_ALIGN_PARAGRAPH.CENTER),
                                       (peca["codigo"],WD_ALIGN_PARAGRAPH.LEFT),
                                       (peca["descricao"],WD_ALIGN_PARAGRAPH.LEFT)]):
            c=tbl_pcs.cell(i+1,j); _bg(c,bg); _bordas(c,cor="CCCCCC")
            _texto_celula(c,val,size=10,align=aln)
    doc.add_paragraph()

    # Observacoes
    p = _par(doc,"OBSERVACOES",bold=True,size=10,depois=2)
    p.runs[0].font.color.rgb = RGBColor(*bytes.fromhex(COR_ESCURA))
    tbl_obs = doc.add_table(rows=1,cols=1)
    tbl_obs.style="Table Grid"; tbl_obs.autofit=False; tbl_obs.columns[0].width=Cm(17)
    c_obs=tbl_obs.cell(0,0); _bg(c_obs,"FFFFFF"); _bordas(c_obs,cor="AAAAAA")
    _texto_celula(c_obs, dados["obs"] if dados["obs"] else "-", size=10)
    doc.add_paragraph()

    # Colaborador
    p = _par(doc,"COLABORADOR QUE EXECUTOU O SERVICO",bold=True,size=10,depois=2)
    p.runs[0].font.color.rgb = RGBColor(*bytes.fromhex(COR_ESCURA))
    tbl_col = doc.add_table(rows=1,cols=1)
    tbl_col.style="Table Grid"; tbl_col.autofit=False; tbl_col.columns[0].width=Cm(17)
    c_col=tbl_col.cell(0,0); _bg(c_col,"FFFFFF"); _bordas(c_col,cor="AAAAAA")
    _texto_celula(c_col, dados["colaborador"], size=10)
    doc.add_paragraph()

    # Assinaturas
    tbl_ass = doc.add_table(rows=2,cols=2)
    tbl_ass.style="Table Grid"; tbl_ass.autofit=False
    tbl_ass.columns[0].width=Cm(8.5); tbl_ass.columns[1].width=Cm(8.5)
    for col,label in enumerate(("Assinatura do Tecnico","Assinatura do Cliente")):
        c_lbl=tbl_ass.cell(0,col); c_sig=tbl_ass.cell(1,col)
        _bg(c_lbl,COR_MEDIA); _bordas(c_lbl,cor="FFFFFF")
        _texto_celula(c_lbl,label,bold=True,size=9,
                      align=WD_ALIGN_PARAGRAPH.CENTER,cor_hex="FFFFFF")
        _bordas(c_sig,cor="AAAAAA"); _texto_celula(c_sig,"",size=16)

    doc.save(destino)


# ─────────────────────────────────────────────
# Widgets auxiliares
# ─────────────────────────────────────────────

def _secao(parent, texto):
    """Faixa de título de seção."""
    frame = tk.Frame(parent, bg=GUI_DARK)
    frame.pack(fill="x", pady=(18, 4))
    tk.Label(frame, text=texto, font=FONT_SECTION,
             bg=GUI_DARK, fg="white", padx=10, pady=4).pack(side="left")
    return frame


def _campo(parent, label, row, col=0, largura=28, obrigatorio=True):
    """Label + Entry; retorna a variável StringVar."""
    tk.Label(parent, text=label + (" *" if obrigatorio else ""),
             font=FONT_LABEL, bg=GUI_CARD, fg=GUI_LABEL,
             anchor="w").grid(row=row, column=col, sticky="w", padx=(0,4), pady=2)
    var = tk.StringVar()
    e = tk.Entry(parent, textvariable=var, font=FONT_ENTRY,
                 width=largura, relief="flat",
                 bg="#EEF2F7", fg="#1A1A2E",
                 insertbackground=GUI_DARK,
                 highlightthickness=1, highlightbackground=GUI_BORDER,
                 highlightcolor=GUI_ACCENT)
    e.grid(row=row, column=col+1, sticky="ew", padx=(0,12), pady=2)
    return var, e


# ─────────────────────────────────────────────
# Linha de peça (widget composto)
# ─────────────────────────────────────────────

class LinhaPeca:
    def __init__(self, parent, idx, on_remove):
        self.frame = tk.Frame(parent, bg=GUI_CARD,
                              highlightthickness=1, highlightbackground=GUI_BORDER)
        self.frame.pack(fill="x", pady=2)

        tk.Label(self.frame, text=str(idx), font=FONT_SMALL,
                 bg=GUI_DARK, fg="white", width=3).pack(side="left")

        self.v_codigo    = tk.StringVar()
        self.v_qtd       = tk.StringVar()
        self.v_descricao = tk.StringVar()

        def _entry(var, ph, w):
            e = tk.Entry(self.frame, textvariable=var, font=FONT_ENTRY,
                         width=w, relief="flat", bg="#EEF2F7", fg="#1A1A2E",
                         insertbackground=GUI_DARK,
                         highlightthickness=1, highlightbackground=GUI_BORDER,
                         highlightcolor=GUI_ACCENT)
            e.pack(side="left", padx=3, pady=4)
            return e

        _entry(self.v_codigo,    "Código",   14)
        _entry(self.v_qtd,       "Qtde.",     5)
        _entry(self.v_descricao, "Descrição", 26)

        tk.Button(self.frame, text="✕", font=FONT_SMALL,
                  bg="#FFEBEE", fg=GUI_ERROR, relief="flat",
                  cursor="hand2", command=lambda: on_remove(self)
                  ).pack(side="left", padx=4)

    def valores(self):
        return (self.v_codigo.get().strip(),
                self.v_qtd.get().strip(),
                self.v_descricao.get().strip())

    def destruir(self):
        self.frame.destroy()


# ─────────────────────────────────────────────
# Aplicação principal
# ─────────────────────────────────────────────

class AppOS:
    def __init__(self, root):
        self.root = root
        root.title("RPM Eletrodiesel — Gerador de O.S.")
        root.configure(bg=GUI_BG)
        root.resizable(True, True)

        self._linhas_peca   = []
        self._linhas_servico = []

        self._build_ui()
        root.update_idletasks()
        w, h = 820, root.winfo_reqheight()
        sw, sh = root.winfo_screenwidth(), root.winfo_screenheight()
        root.geometry("{}x{}+{}+{}".format(
            w, min(h+40, sh-60),
            (sw-w)//2, max(20,(sh-h)//2)))

    # ── Construção da interface ───────────────

    def _build_ui(self):
        # Barra superior
        bar = tk.Frame(self.root, bg=GUI_DARK, height=52)
        bar.pack(fill="x")
        bar.pack_propagate(False)
        tk.Label(bar, text="RPM ELETRODIESEL",
                 font=("Segoe UI", 13, "bold"),
                 bg=GUI_DARK, fg="white").pack(side="left", padx=16, pady=12)
        tk.Label(bar, text="Gerador de Ordem de Serviço",
                 font=("Segoe UI", 10),
                 bg=GUI_DARK, fg="#8AABCF").pack(side="left", pady=12)

        # Área rolável
        canvas = tk.Canvas(self.root, bg=GUI_BG, highlightthickness=0)
        sb = ttk.Scrollbar(self.root, orient="vertical", command=canvas.yview)
        canvas.configure(yscrollcommand=sb.set)
        sb.pack(side="right", fill="y")
        canvas.pack(side="left", fill="both", expand=True)

        self.inner = tk.Frame(canvas, bg=GUI_BG)
        win_id = canvas.create_window((0, 0), window=self.inner, anchor="nw")

        def _resize(e):
            canvas.configure(scrollregion=canvas.bbox("all"))
            canvas.itemconfig(win_id, width=e.width)

        self.inner.bind("<Configure>", _resize)
        canvas.bind("<Configure>", lambda e: canvas.itemconfig(win_id, width=e.width))

        # Scroll com roda do mouse
        def _scroll(e):
            canvas.yview_scroll(int(-1*(e.delta/120)), "units")
        self.root.bind_all("<MouseWheel>", _scroll)

        pad = tk.Frame(self.inner, bg=GUI_BG)
        pad.pack(fill="both", expand=True, padx=24, pady=16)

        self._card_os(pad)
        self._card_cliente(pad)
        self._card_servicos(pad)
        self._card_pecas(pad)
        self._card_obs_colab(pad)
        self._botao_gerar(pad)

    def _card(self, parent):
        f = tk.Frame(parent, bg=GUI_CARD,
                     highlightthickness=1, highlightbackground=GUI_BORDER)
        f.pack(fill="x", pady=(0, 10))
        inner = tk.Frame(f, bg=GUI_CARD)
        inner.pack(fill="x", padx=16, pady=12)
        return inner

    # ── Seções ────────────────────────────────

    def _card_os(self, parent):
        _secao(parent, "  DADOS DA ORDEM DE SERVIÇO")
        c = self._card(parent)
        c.columnconfigure(1, weight=1)
        c.columnconfigure(3, weight=1)
        self.v_os_num, _ = _campo(c, "Número da O.S.", 0, col=0, largura=20)
        self.v_data,   _ = _campo(c, "Data (DD/MM/AAAA)", 0, col=2, largura=16)

    def _card_cliente(self, parent):
        _secao(parent, "  DADOS DO CLIENTE / VEÍCULO")
        c = self._card(parent)
        c.columnconfigure(1, weight=2)
        c.columnconfigure(3, weight=1)
        self.v_cliente, _ = _campo(c, "Cliente",  0, col=0, largura=32)
        self.v_cidade,  _ = _campo(c, "Cidade",   0, col=2, largura=22)
        self.v_motor,   _ = _campo(c, "Motor",    1, col=0, largura=32)
        self.v_veiculo, _ = _campo(c, "Veículo",  1, col=2, largura=22)
        self.v_placa,   _ = _campo(c, "Placa",    2, col=0, largura=12)
        tk.Label(c, text="máx. 8 caracteres", font=FONT_SMALL,
                 bg=GUI_CARD, fg=GUI_LABEL).grid(row=2, column=2, sticky="w")

    def _card_servicos(self, parent):
        _secao(parent, "  SERVIÇOS REALIZADOS")
        c = self._card(parent)
        self._frame_servicos = tk.Frame(c, bg=GUI_CARD)
        self._frame_servicos.pack(fill="x")
        tk.Button(c, text="+ Adicionar Serviço",
                  font=FONT_SMALL, bg=GUI_ACCENT, fg="white",
                  relief="flat", cursor="hand2", padx=8, pady=4,
                  command=self._add_servico).pack(anchor="w", pady=(6,0))
        self._add_servico()   # começa com 1 linha

    def _add_servico(self):
        idx = len(self._linhas_servico) + 1
        row = tk.Frame(self._frame_servicos, bg=GUI_CARD,
                       highlightthickness=1, highlightbackground=GUI_BORDER)
        row.pack(fill="x", pady=2)
        tk.Label(row, text=str(idx), font=FONT_SMALL,
                 bg=GUI_DARK, fg="white", width=3).pack(side="left")
        var = tk.StringVar()
        tk.Entry(row, textvariable=var, font=FONT_ENTRY,
                 relief="flat", bg="#EEF2F7", fg="#1A1A2E",
                 insertbackground=GUI_DARK,
                 highlightthickness=1, highlightbackground=GUI_BORDER,
                 highlightcolor=GUI_ACCENT).pack(
                     side="left", fill="x", expand=True, padx=4, pady=4)

        def _rem(r=row, v=var):
            if len(self._linhas_servico) <= 1:
                messagebox.showwarning("Atenção", "É necessário ao menos 1 serviço.")
                return
            self._linhas_servico.remove((r, v)); r.destroy()

        tk.Button(row, text="✕", font=FONT_SMALL,
                  bg="#FFEBEE", fg=GUI_ERROR, relief="flat",
                  cursor="hand2", command=_rem).pack(side="left", padx=4)
        self._linhas_servico.append((row, var))

    def _card_pecas(self, parent):
        _secao(parent, "  PEÇAS UTILIZADAS")
        c = self._card(parent)

        # Cabeçalho das colunas
        hdr = tk.Frame(c, bg=GUI_DARK)
        hdr.pack(fill="x", pady=(0,4))
        for txt, w in (("Nº",3),("Código",14),("Qtde.",5),("Descrição",26),("",4)):
            tk.Label(hdr, text=txt, font=FONT_SMALL,
                     bg=GUI_DARK, fg="white", width=w,
                     anchor="w", padx=4).pack(side="left")

        self._frame_pecas = tk.Frame(c, bg=GUI_CARD)
        self._frame_pecas.pack(fill="x")

        tk.Button(c, text="+ Adicionar Peça",
                  font=FONT_SMALL, bg=GUI_ACCENT, fg="white",
                  relief="flat", cursor="hand2", padx=8, pady=4,
                  command=self._add_peca).pack(anchor="w", pady=(6,0))
        self._add_peca()   # começa com 1 linha

    def _add_peca(self):
        idx = len(self._linhas_peca) + 1
        lp = LinhaPeca(self._frame_pecas, idx, self._rem_peca)
        self._linhas_peca.append(lp)

    def _rem_peca(self, lp):
        if len(self._linhas_peca) <= 1:
            messagebox.showwarning("Atenção", "É necessário ao menos 1 peça.")
            return
        self._linhas_peca.remove(lp)
        lp.destruir()

    def _card_obs_colab(self, parent):
        _secao(parent, "  OBSERVAÇÕES E EXECUÇÃO")
        c = self._card(parent)

        tk.Label(c, text="Observações", font=FONT_LABEL,
                 bg=GUI_CARD, fg=GUI_LABEL).pack(anchor="w")
        self.txt_obs = tk.Text(c, font=FONT_ENTRY, height=3,
                               relief="flat", bg="#EEF2F7", fg="#1A1A2E",
                               insertbackground=GUI_DARK,
                               highlightthickness=1, highlightbackground=GUI_BORDER,
                               highlightcolor=GUI_ACCENT)
        self.txt_obs.pack(fill="x", pady=(2,10))

        tk.Label(c, text="Colaborador que executou o serviço *",
                 font=FONT_LABEL, bg=GUI_CARD, fg=GUI_LABEL).pack(anchor="w")
        self.v_colaborador = tk.StringVar()
        tk.Entry(c, textvariable=self.v_colaborador, font=FONT_ENTRY,
                 relief="flat", bg="#EEF2F7", fg="#1A1A2E",
                 insertbackground=GUI_DARK,
                 highlightthickness=1, highlightbackground=GUI_BORDER,
                 highlightcolor=GUI_ACCENT).pack(fill="x", pady=(2,0))

    def _botao_gerar(self, parent):
        tk.Label(parent, text="* Campos obrigatórios",
                 font=FONT_SMALL, bg=GUI_BG, fg=GUI_LABEL).pack(anchor="w")

        self.lbl_status = tk.Label(parent, text="", font=FONT_LABEL, bg=GUI_BG)
        self.lbl_status.pack(anchor="w", pady=(4,0))

        tk.Button(parent, text="  GERAR ORDEM DE SERVIÇO  ",
                  font=FONT_BTN, bg=GUI_DARK, fg=GUI_BTN_FG,
                  relief="flat", cursor="hand2",
                  activebackground=GUI_ACCENT, activeforeground="white",
                  padx=20, pady=12,
                  command=self._gerar).pack(anchor="e", pady=16)

    # ── Validação e geração ───────────────────

    def _status(self, msg, cor=GUI_ERROR):
        self.lbl_status.config(text=msg, fg=cor)
        self.root.update_idletasks()

    def _gerar(self):
        self._status("")

        # ── Coleta e valida ──
        erros = []

        os_num = self.v_os_num.get().strip()
        if not os_num: erros.append("Número da O.S.")

        data = self.v_data.get().strip()
        if not data:
            data = datetime.today().strftime("%d/%m/%Y")
        else:
            padrao = re.compile(r"^\d{2}/\d{2}/\d{4}$")
            if not padrao.match(data):
                erros.append("Data (formato inválido — use DD/MM/AAAA)")
            else:
                try:
                    d,m,a = data.split("/")
                    datetime(int(a),int(m),int(d))
                except ValueError:
                    erros.append("Data (data inexistente)")

        cliente = self.v_cliente.get().strip()
        if not cliente: erros.append("Cliente")

        cidade = self.v_cidade.get().strip()
        if not cidade: erros.append("Cidade")

        motor = self.v_motor.get().strip()
        if not motor: erros.append("Motor")

        veiculo = self.v_veiculo.get().strip()
        if not veiculo: erros.append("Veículo")

        placa = self.v_placa.get().strip().upper()
        if placa and len(placa) > 8:
            erros.append("Placa ({} caracteres — máximo 8)".format(len(placa)))

        # Serviços
        servicos = []
        for i, (_, var) in enumerate(self._linhas_servico):
            s = var.get().strip()
            if i == 0 and not s:
                erros.append("Serviço 01 (obrigatório)")
            elif s:
                servicos.append(s)
        if not servicos and not any("Serviço 01" in e for e in erros):
            erros.append("Ao menos 1 serviço é obrigatório")

        # Peças
        pecas = []
        for i, lp in enumerate(self._linhas_peca):
            cod, qtd_str, desc = lp.valores()
            if i == 0 and not cod:
                erros.append("Código da 1ª peça (obrigatório)")
                continue
            if not cod: continue   # linha vazia após a 1ª = ignorar
            if not qtd_str or not qtd_str.isdigit() or int(qtd_str) < 1:
                erros.append("Quantidade da peça {} (número inteiro > 0)".format(i+1))
                continue
            if not desc:
                erros.append("Descrição da peça {} (obrigatória)".format(i+1))
                continue
            pecas.append({"codigo": cod, "qtd": int(qtd_str), "descricao": desc})
        if not pecas and not any("peça" in e for e in erros):
            erros.append("Ao menos 1 peça é obrigatória")

        obs         = self.txt_obs.get("1.0", "end").strip()
        colaborador = self.v_colaborador.get().strip()
        if not colaborador: erros.append("Colaborador que executou o serviço")

        if erros:
            self._status("Corrija os campos: " + " | ".join(erros))
            return

        # ── Escolha do destino ──
        cliente_safe = re.sub(r'[\\/:*?"<>|]', "", cliente).strip()
        nome_sugerido = "OS_{}_{}_{}.docx".format(
            os_num, cliente_safe, data.replace("/",""))
        destino = filedialog.asksaveasfilename(
            title="Salvar Ordem de Serviço",
            initialfile=nome_sugerido,
            defaultextension=".docx",
            filetypes=[("Word", "*.docx"), ("Todos", "*.*")])
        if not destino:
            return

        # ── Gera ──
        self._status("Gerando documento...", cor=GUI_ACCENT)
        try:
            dados = {
                "os_numero":   os_num,
                "data_os":     data,
                "cliente":     cliente,
                "cidade":      cidade,
                "motor":       motor,
                "veiculo":     veiculo,
                "placa":       placa,
                "servicos":    servicos,
                "pecas":       pecas,
                "obs":         obs,
                "colaborador": colaborador,
            }
            gerar_os(dados, destino)
            self._status("✔  Arquivo salvo em: " + destino, cor=GUI_SUCCESS)
            messagebox.showinfo("Sucesso",
                "Ordem de Serviço gerada com sucesso!\n\n" + destino)
        except Exception as ex:
            self._status("Erro ao gerar documento: " + str(ex))
            messagebox.showerror("Erro", str(ex))


# ─────────────────────────────────────────────
# Ponto de entrada
# ─────────────────────────────────────────────

if __name__ == "__main__":
    root = tk.Tk()
    AppOS(root)
    root.mainloop()